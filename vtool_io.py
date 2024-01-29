import string
import sys
import os
import time
from antlr4 import FileStream, CommonTokenStream

from verilog.VerilogLexer import VerilogLexer
from verilog.VerilogParser import VerilogParser
from verilog.VerilogParserVisitor import VerilogParserVisitor
from verilog.VerilogPreParser import VerilogPreParser
from verilog.VerilogPreParserVisitor import VerilogPreParserVisitor

import re

from docx import Document
from docx.shared import RGBColor
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL

class ParseResult:
  def __init__(self):
    self.modules = []
    self.ports = [] # list of dict
    self.params = []

  def set_ports(self, matches):
    for i, match in enumerate(matches):
      dir, type, start, end, name = match

      result = {}
      if dir    != '': result['dir'] = dir
      if type   != '': result['type'] = type
      if start  != '': 
        result.update({'msb': start, 'lsb': end, 'name': name })
      else : 
        result['name'] = name

      self.ports.append(result)
      #print( "name: ", name, result)



  def get_port_matches(self, text):
    #print(self.params)

    #replace param
    for i, param in enumerate(self.params):
      param_name = param['name']
      param_value = param['value']
      param_value = int(param_value) - 1
      param_value = str(param_value)
      pattern = f'{param_name}-1'
      #print (pattern)
      #text = text.replace(pattern, param_value)
      text = re.sub(pattern, param_value, text)

    #print(text)

    matches = re.findall(r'(input|output|inout)(wire|reg|signed)?(?:\[(\d+):(\d+)\])?(\w+)', text)

    if len(matches) > 0:
      self.set_ports(matches)
    else:
      print("No results to update.")

  def get_param_matches(self, text):
    matches = re.findall(r'(\w+)\s*=\s*(\d+)', text)

    #print (matches)

    for i, match in enumerate(matches):
      result = {}
      name, value = match

      name = name.replace('parameter', '')
      result['name'] = name
      result['value'] = value

      self.params.append(result)

    #print (self.params)

  def gen_io(self):
    # 创建Word文档对象
    doc = Document()

    # 创建表格
    table = doc.add_table(rows=(len(self.ports)+1), cols=4)

    # 设置表格样式
    table.style = 'Table Grid'

    # 设置标题栏底色为蓝色
    title_row = table.rows[0]
    for cell in title_row.cells:
      shading = cell._element.xpath('.//w:shd')
      if shading:
        shading[0].attrib['fill'] = '0000FF'  # 蓝色

    # 设置表格线样式
    for row in table.rows:
      for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.vertical_anchor = WD_ALIGN_VERTICAL.CENTER

        for paragraph in cell.paragraphs:
          for run in paragraph.runs:
            run.font.size = 12  # 设置字体大小

    # 填充表格内容
    table.cell(0, 0).text = 'Name'
    table.cell(0, 1).text = 'Width'
    table.cell(0, 2).text = 'Dir'
    table.cell(0, 3).text = 'Desc'

    #for port in self.ports:
    #  print("Signal: ", port)

    last_dir = ''
    cur_dir = ''
    add_row_id = 1

    for i, port in enumerate(self.ports):
      cur_dir = port['dir']

      if last_dir != cur_dir and i != 0:
        # Insert a blank row and merge its cells
        table.add_row()
        table.cell(i+add_row_id, 0).merge(table.cell(i+add_row_id, 3))
        add_row_id += 1

      table.cell(i+add_row_id, 0).text = port['name']

      if 'msb' in port and port['msb'] != '':
        table.cell(i+add_row_id, 1).text = str(int(port['msb']) - int(port['lsb']) + 1)
      else:
        table.cell(i+add_row_id, 1).text = '1'



      if port['dir'] == 'input':
        table.cell(i+add_row_id, 2).text = "I"
      elif port['dir'] == 'output':
        table.cell(i+add_row_id, 2).text = "O"

      table.cell(i+add_row_id, 3).text = ''

      last_dir = cur_dir

    doc.save('table.docx')
      
class PortVisitor(VerilogParserVisitor):
  def __init__(self, results):
    self.results = results

  def visitPort_declaration(self, ctx: VerilogParser.Port_declarationContext):
    rtn = ctx.inout_declaration()
    if rtn is not None: 
      #print('inout:',rtn.getText())
      self.results.get_port_matches(rtn.getText())
 
    rtn = ctx.input_declaration()
    if rtn is not None: 
      #print('input:',rtn.getText())
      self.results.get_port_matches(rtn.getText())

    rtn = ctx.output_declaration()
    if rtn is not None: 
      #print('output:',rtn.getText())
      self.results.get_port_matches(rtn.getText())

class ModuleVisitor(VerilogParserVisitor):
  def __init__(self, results):
    self.results = results

  def visitModule_declaration(self, ctx:VerilogParser.Module_declarationContext):
    rtn = ctx.module_identifier()
    if rtn is not None: 
      #print('module_name:',rtn.getText())
      self.results.modules.append(rtn.getText())

    rtn = ctx.module_parameter_port_list()
    if rtn is not None: 
      #print('param_list:',rtn.getText())
      self.results.get_param_matches(rtn.getText())
      
    #rtn = ctx.list_of_port_declarations()
    #if rtn is not None: 
    #  print('port_list:',rtn.getText())
    #  #self.results.ports.append(rtn.getText())
    #  #self.results.get_port_matches(rtn.getText())


def main(argv):
  start = time.time()

  if len(sys.argv) > 1:
      filename = sys.argv[1]
  else:
      filename = 'test.v'  # default filename

  if not os.path.exists(filename):
      print(f"File {filename} does not exist.")
      sys.exit(1)

  input_stream = FileStream(filename)
  lexer = VerilogLexer(input_stream)
  stream = CommonTokenStream(lexer)
  parser = VerilogParser(stream)
  context = parser.source_text()

  parse_result = ParseResult()

  visitor_module = ModuleVisitor(parse_result)
  ast = visitor_module.visitSource_text(context)

  visitor_port = PortVisitor(parse_result)
  ast = visitor_port.visitSource_text(context)
  
  #print (ast)

  #print("=====================================")
  #for port in parse_result.ports:
  #  print("Signal: ", port)

  parse_result.gen_io()

  print(f"Elapsed time: {time.time() - start} s")

if __name__ == "__main__":
  main(sys.argv)